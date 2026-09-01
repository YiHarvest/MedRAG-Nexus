"""校验并解析受支持的文件，将内容统一转换为 Markdown。"""

from __future__ import annotations

import asyncio
import codecs
import re
import tempfile
import time
import zipfile
from collections.abc import Awaitable, Callable
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit, urlunsplit

import httpx

from medrag_nexus.core.config import Settings

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ParseProgress = Callable[[str, str, dict[str, object]], Awaitable[None]]


@dataclass(slots=True)
class ParseResult:
    markdown: str
    parser: str
    degraded: bool = False


async def _report(
    progress: ParseProgress | None,
    level: str,
    message: str,
    **context: object,
) -> None:
    if progress is not None:
        await progress(level, message, context)


def _safe_endpoint(url: str) -> str:
    """保留服务定位信息，同时移除 URL 中可能携带的账号、查询参数和片段。"""
    parsed = urlsplit(url)
    host = parsed.hostname or ""
    if parsed.port:
        host = f"{host}:{parsed.port}"
    return urlunsplit((parsed.scheme, host, parsed.path, "", ""))


def safe_file_name(name: str) -> str:
    value = Path(name).name
    value = re.sub(r"[^\w.() -]+", "_", value, flags=re.UNICODE).strip(" .")
    if not value or value in {".", ".."}:
        raise ValueError("invalid file name")
    return value[:255]


def sniff_extension(path: Path) -> str:
    head = path.read_bytes()[:16]
    if head.startswith(b"%PDF-"):
        return ".pdf"
    if zipfile.is_zipfile(path):
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
        if "word/document.xml" in names:
            return ".docx"
    try:
        decoder = codecs.getincrementaldecoder("utf-8")()
        decoder.decode(head, final=False)
    except UnicodeDecodeError:
        return ""
    return ".txt"


def validate_file_type(path: Path) -> str:
    declared = path.suffix.lower()
    if declared not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported file extension: {declared or '<none>'}")
    detected = sniff_extension(path)
    if not detected or declared != detected:
        detected_label = detected or "unknown"
        raise ValueError(f"file content does not match extension: declared={declared}, detected={detected_label}")
    return declared


def extract_mineru_markdown(payload: Any) -> str:
    if not isinstance(payload, dict):
        return ""
    candidates = [payload.get("markdown"), payload.get("md"), payload.get("content")]
    data = payload.get("data")
    if isinstance(data, dict):
        candidates.extend([data.get("markdown"), data.get("md_content")])
    results = payload.get("results")
    if isinstance(results, dict):
        for result in results.values():
            if isinstance(result, dict):
                candidates.extend([result.get("md_content"), result.get("markdown"), result.get("md")])
    return next((item.strip() for item in candidates if isinstance(item, str) and item.strip()), "")


def _read_mineru_markdown(output_dir: Path) -> str:
    markdown_files = sorted(output_dir.rglob("*.md"))
    for markdown_file in markdown_files:
        markdown = markdown_file.read_text(encoding="utf-8").strip()
        if markdown:
            return markdown
    raise RuntimeError("MinerU returned no markdown")


async def _mineru_http_client(
    path: Path,
    settings: Settings,
    progress: ParseProgress | None = None,
) -> str:
    from mineru.cli.common import aio_do_parse, read_fn
    from mineru.utils.enum_class import MakeMode

    with tempfile.TemporaryDirectory(prefix="medrag-nexus-mineru-") as temporary_dir:
        output_dir = Path(temporary_dir)
        started = time.monotonic()
        size_bytes = (await asyncio.to_thread(path.stat)).st_size
        await _report(
            progress,
            "INFO",
            "正在连接 MinerU OpenAI 服务并提交文件",
            endpoint=_safe_endpoint(settings.mineru_url),
            backend=settings.mineru_backend,
            method=settings.mineru_method,
            language=settings.mineru_lang,
            size_bytes=size_bytes,
        )
        await aio_do_parse(
            output_dir=str(output_dir),
            pdf_file_names=[path.stem],
            pdf_bytes_list=[read_fn(path)],
            p_lang_list=[settings.mineru_lang],
            backend=settings.mineru_backend,
            parse_method=settings.mineru_method,
            formula_enable=True,
            table_enable=True,
            server_url=settings.mineru_url.rstrip("/"),
            f_draw_layout_bbox=False,
            f_draw_span_bbox=False,
            f_dump_md=True,
            f_dump_middle_json=False,
            f_dump_model_output=False,
            f_dump_orig_pdf=False,
            f_dump_content_list=False,
            f_make_md_mode=MakeMode.NLP_MD,
            image_analysis=False,
            max_concurrency=settings.mineru_max_concurrency,
            http_timeout=settings.mineru_http_timeout_seconds,
        )
        await _report(
            progress,
            "INFO",
            "MinerU OpenAI 服务响应完成，开始读取解析产物",
            elapsed_ms=round((time.monotonic() - started) * 1000),
        )
        markdown = await asyncio.to_thread(_read_mineru_markdown, output_dir)
        await _report(progress, "INFO", "MinerU 解析产物读取完成", characters=len(markdown))
        return markdown


async def _mineru_file_api(
    path: Path,
    settings: Settings,
    progress: ParseProgress | None = None,
) -> str:
    if not settings.mineru_url:
        raise RuntimeError("MinerU is not configured")
    fields = {
        "backend": settings.mineru_backend,
        "parse_method": settings.mineru_method,
        "lang_list": settings.mineru_lang,
        "formula_enable": "true",
        "table_enable": "true",
        "return_md": "true",
        "return_middle_json": "false",
        "return_model_output": "false",
        "return_content_list": "false",
        "return_images": "false",
        "response_format_zip": "false",
    }
    endpoint = settings.mineru_url.rstrip("/")
    if not endpoint.endswith(settings.mineru_api_path):
        endpoint = f"{endpoint}/{settings.mineru_api_path.lstrip('/')}"
    size_bytes = (await asyncio.to_thread(path.stat)).st_size
    timeout_seconds = getattr(settings, "mineru_http_timeout_seconds", 1800)
    
    max_retries = 3
    retry_delay = 10.0
    last_error: Exception | None = None
    
    for attempt in range(max_retries):
        try:
            await _report(
                progress,
                "INFO",
                "准备连接 MinerU 文件解析接口",
                endpoint=_safe_endpoint(endpoint),
                backend=settings.mineru_backend,
                method=settings.mineru_method,
                language=settings.mineru_lang,
                timeout_seconds=timeout_seconds,
                size_bytes=size_bytes,
                attempt=attempt + 1,
                max_retries=max_retries,
            )
            started = time.monotonic()
            
            async with httpx.AsyncClient(
                timeout=httpx.Timeout(
                    connect=30.0,
                    read=timeout_seconds,
                    write=300.0,
                    pool=30.0,
                ),
                follow_redirects=False,
                trust_env=False,
                limits=httpx.Limits(max_keepalive_connections=1, max_connections=1),
            ) as client:
                with path.open("rb") as stream:
                    await _report(progress, "INFO", "开始向 MinerU 上传并解析文件", endpoint=_safe_endpoint(endpoint))
                    response = await client.post(
                        endpoint,
                        data=fields,
                        files={"files": (path.name, stream, "application/octet-stream")},
                    )
                await _report(
                    progress,
                    "INFO" if response.is_success else "ERROR",
                    "MinerU 接口响应已收到",
                    endpoint=_safe_endpoint(endpoint),
                    status=response.status_code,
                    content_type=response.headers.get("content-type", "unknown"),
                    elapsed_ms=round((time.monotonic() - started) * 1000),
                )
                response.raise_for_status()
            
            if response.headers.get("content-type", "").startswith("application/json"):
                payload: Any = response.json()
                markdown = extract_mineru_markdown(payload)
            else:
                markdown = response.text
            if not markdown.strip():
                raise RuntimeError("MinerU returned no markdown")
            markdown = markdown.strip()
            await _report(progress, "INFO", "MinerU Markdown 解析完成", characters=len(markdown))
            return markdown
            
        except (httpx.ConnectError, httpx.ReadError, httpx.WriteError, httpx.TimeoutException) as error:
            last_error = error
            error_type = type(error).__name__
            await _report(
                progress,
                "WARNING",
                "MinerU 连接异常，准备重试",
                endpoint=_safe_endpoint(endpoint),
                error_type=error_type,
                error_message=str(error),
                attempt=attempt + 1,
                max_retries=max_retries,
                retry_delay_seconds=retry_delay,
            )
            if attempt < max_retries - 1:
                await asyncio.sleep(retry_delay)
                retry_delay = min(retry_delay * 2, 60.0)
            else:
                await _report(
                    progress,
                    "ERROR",
                    "MinerU 连接失败，已达最大重试次数",
                    endpoint=_safe_endpoint(endpoint),
                    error_type=error_type,
                    attempts=max_retries,
                )
                raise RuntimeError(f"MinerU connection failed after {max_retries} retries: {error_type}") from error
                
        except httpx.HTTPStatusError as error:
            await _report(
                progress,
                "ERROR",
                "MinerU 返回错误状态码",
                endpoint=_safe_endpoint(endpoint),
                status_code=error.response.status_code,
                error_message=str(error),
            )
            raise
            
        except Exception as error:
            last_error = error
            error_type = type(error).__name__
            await _report(
                progress,
                "ERROR",
                "MinerU 处理异常",
                endpoint=_safe_endpoint(endpoint),
                error_type=error_type,
                error_message=str(error),
            )
            raise
    
    if last_error:
        raise last_error
    raise RuntimeError("MinerU processing failed")


async def _mineru(path: Path, settings: Settings, progress: ParseProgress | None = None) -> str:
    if not settings.mineru_url:
        raise RuntimeError("MinerU is not configured")
    if settings.mineru_uses_openai_server:
        return await _mineru_http_client(path, settings, progress)
    return await _mineru_file_api(path, settings, progress)


def _pdf(path: Path) -> str:
    from pypdf import PdfReader

    pages = []
    for number, page in enumerate(PdfReader(str(path)).pages, start=1):
        pages.append(f"<!-- page: {number} -->\n\n{page.extract_text() or ''}")
    return "\n\n".join(pages)


def _docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    rows: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name.lower() if paragraph.style else ""
        if style.startswith("heading"):
            level = re.search(r"\d+", style)
            rows.append(f"{'#' * min(int(level.group()) if level else 2, 6)} {text}")
        else:
            rows.append(text)
    for table in document.tables:
        values = [[cell.text.strip().replace("|", "\\|") for cell in row.cells] for row in table.rows]
        if values:
            rows.append(_markdown_table(values))
    return "\n\n".join(rows)


def _markdown_table(rows: list[list[str]]) -> str:
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    return "\n".join(
        ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
        + ["| " + " | ".join(row) + " |" for row in padded[1:]]
    )


def _text(path: Path) -> str:
    from charset_normalizer import from_bytes

    match = from_bytes(path.read_bytes()).best()
    if match is None:
        raise ValueError("unable to detect text encoding")
    return str(match)


_FALLBACKS = {
    ".pdf": ("pypdf", _pdf),
    ".docx": ("python-docx", _docx),
}


async def parse_file(
    path: Path,
    settings: Settings,
    *,
    progress: ParseProgress | None = None,
) -> ParseResult:
    started = time.monotonic()
    extension = validate_file_type(path)
    size_bytes = (await asyncio.to_thread(path.stat)).st_size
    await _report(
        progress,
        "INFO",
        "文件类型校验完成",
        extension=extension,
        size_bytes=size_bytes,
    )
    if extension == ".txt":
        await _report(progress, "INFO", "开始检测文本编码并提取内容", parser="text")
        markdown = await asyncio.to_thread(_text, path)
        await _report(
            progress,
            "INFO",
            "文本文件解析完成",
            parser="text",
            characters=len(markdown),
            elapsed_ms=round((time.monotonic() - started) * 1000),
        )
        return ParseResult(markdown=markdown, parser="text")
    await _report(
        progress,
        "INFO",
        "开始调用 MinerU 解析文件",
        parser="mineru",
        backend=settings.mineru_backend,
        method=settings.mineru_method,
    )
    try:
        markdown = await _mineru(path, settings, progress)
        await _report(
            progress,
            "INFO",
            "MinerU 文件解析完成",
            parser="mineru",
            characters=len(markdown),
            elapsed_ms=round((time.monotonic() - started) * 1000),
        )
        return ParseResult(markdown=markdown, parser="mineru")
    except Exception as mineru_error:
        fallback = _FALLBACKS.get(extension)
        if fallback is None:
            message = f"MinerU failed and no fallback exists for {extension}: {mineru_error}"
            raise RuntimeError(message) from mineru_error
        parser_name, parser = fallback
        await _report(
            progress,
            "WARN",
            "MinerU 解析失败，切换本地解析器",
            parser="mineru",
            fallback_parser=parser_name,
            exception_type=type(mineru_error).__name__,
            error=str(mineru_error)[:500],
        )
        fallback_started = time.monotonic()
        markdown = await asyncio.to_thread(parser, path)
        if not markdown.strip():
            raise RuntimeError(f"{parser_name} returned no content after MinerU failed") from mineru_error
        await _report(
            progress,
            "WARN",
            "本地降级解析完成",
            parser=parser_name,
            characters=len(markdown),
            elapsed_ms=round((time.monotonic() - fallback_started) * 1000),
        )
        return ParseResult(markdown=markdown, parser=parser_name, degraded=True)
