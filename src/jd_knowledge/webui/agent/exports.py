"""在内存中把 Agent 回答渲染为 Word。"""

from __future__ import annotations

import asyncio
from datetime import datetime, timedelta
from io import BytesIO

from docx import Document
from docx.shared import Pt

from .artifacts import ArtifactService
from .models import AgentArtifact, AnswerExportContent, ArtifactResourceRequirement
from .store import utc_now

WORD_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def render_answer_docx(content: AnswerExportContent, *, now: datetime | None = None) -> bytes:
    """完全在内存中渲染回答，源文本不会另行持久化。"""

    generated_at = content.generated_at or now or utc_now()
    document = Document()
    document.core_properties.title = content.title
    document.add_heading(content.title, level=0)
    metadata = f"生成时间：{generated_at.astimezone().strftime('%Y-%m-%d %H:%M:%S %Z')}"
    if content.generated_by:
        metadata += f"\n生成账号：{content.generated_by}"
    paragraph = document.add_paragraph(metadata)
    paragraph.style.font.size = Pt(9)
    document.add_heading("用户问题", level=1)
    document.add_paragraph(content.question)
    document.add_heading("助手回答", level=1)
    for block in content.answer.split("\n"):
        document.add_paragraph(block)
    if content.sources:
        document.add_heading("引用来源", level=1)
        for index, source in enumerate(content.sources, start=1):
            line = f"{index}. {source.title}"
            if source.reference:
                line += f" — {source.reference}"
            document.add_paragraph(line)
            if source.excerpt:
                document.add_paragraph(source.excerpt, style="Quote")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


async def export_answer_to_word(
    artifact_service: ArtifactService,
    *,
    owner_account_id: str,
    content: AnswerExportContent,
    file_name: str = "知识助手回答.docx",
    conversation_id: str | None = None,
    message_id: str | None = None,
    required_permissions: tuple[str, ...] | list[str] = (),
    resource_requirements: tuple[ArtifactResourceRequirement, ...] | list[ArtifactResourceRequirement] = (),
    ttl: timedelta = ArtifactService.DEFAULT_TTL,
    now: datetime | None = None,
) -> AgentArtifact:
    rendered = await asyncio.to_thread(render_answer_docx, content, now=now)
    return await artifact_service.create(
        owner_account_id=owner_account_id,
        file_name=file_name,
        mime_type=WORD_MIME_TYPE,
        content=rendered,
        conversation_id=conversation_id,
        message_id=message_id,
        required_permissions=required_permissions,
        resource_requirements=resource_requirements,
        ttl=ttl,
        now=now,
    )
